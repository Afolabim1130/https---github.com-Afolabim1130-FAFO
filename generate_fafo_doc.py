from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create document
doc = Document()

def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    return p

# Title
add_heading('FAFO - Full Project Documentation', level=0)
add_paragraph('Generated documentation for the FAFO (Forensic Analysis & Forensic Operations) project.')

# 1. System design explanation
add_heading('1. System Design Explanation', level=1)
add_paragraph('Overview:\nFAFO is a Streamlit-based web application designed to collect, store, and manage multimedia evidence for online incidents. The architecture is modular, combining a Streamlit frontend with a Python backend composed of several modules for authentication, OCR, video processing, hashing, storage management, notifications, and role-based access control.')
add_paragraph('\nComponents:\n- Frontend: `app.py` built with Streamlit.\n- Modules: `modules/` directory containing business logic (auth, ocr_engine, ffmpeg_processor, repository, etc.).\n- Configuration: `config/` with `settings.py` and `roles.py`.\n- Storage: `evidence_repository/` for uploaded evidence and a `database/` folder containing schema and initialization scripts.\n- External binaries: Tesseract (OCR) and FFmpeg/ffprobe (video processing) detected at runtime with fallbacks.')
add_paragraph('\nData Flow:\n1. User logs in and selects a role (submitter/reviewer/lawyer/admin).\n2. Submitter fills form and uploads evidence or provides source URL.\n3. Backend generates Incident ID, timestamps, stores files into an organized incident folder, computes hashes, runs OCR or media analysis, and logs activity.\n4. Notifications are sent automatically to configured recipients.\n5. Reviewers and admins manage evidence and approve access for lawyers.')

# 2. Screenshots of the website (Placeholders)
add_heading('2. Screenshots of the Website', level=1)
add_paragraph('NOTE: Replace the placeholders below with actual screenshots taken from the running app. Place images in the `docs/screenshots/` folder and re-generate this document or insert them manually in Word.')
add_paragraph('Screenshot: Login Page - [PLACEHOLDER]')
add_paragraph('Screenshot: Submitter Form - [PLACEHOLDER]')
add_paragraph('Screenshot: Reviewer Dashboard - [PLACEHOLDER]')
add_paragraph('Screenshot: Lawyer Portal - [PLACEHOLDER]')

# 3. Database or storage explanation
add_heading('3. Database / Storage Explanation', level=1)
add_paragraph('Storage Model:\nFAFO supports a filesystem-based evidence repository and an optional relational database for metadata. The `database/schema.sql` defines the tables for incidents, evidence items, users, roles, and audit trails. The recommended production deployment uses a managed database (PostgreSQL) for robustness and the filesystem (or object storage like S3) for large binary evidence files.')
add_paragraph('\nKey tables and purpose:\n- incidents: Incident metadata (id, title, category, status, timestamps).\n- evidence: Records for each evidence file (path, hash, mimetype, size, incident_id, uploaded_by).\n- users: Authentication and role assignment.\n- audit_logs: Immutable log of actions for chain-of-custody and compliance.')
add_paragraph('\nStorage best practices:\n- Store binary evidence separately from metadata.\n- Use object storage with versioning and lifecycle rules for retention.\n- Keep cryptographic hashes (SHA-256) in the database for integrity verification.')

# 4. Evidence folder structure
add_heading('4. Evidence Folder Structure', level=1)
add_paragraph('Recommended directory layout for each incident:')

p = doc.add_paragraph()
p.add_run('- evidence_repository/').bold = True
p.add_run('\n    - INC-<YYYYMMDD>-<ID>/')
p.add_run('\n        - documents/    (PDFs, text reports)')
p.add_run('\n        - images/       (screenshots, photos)')
p.add_run('\n        - audio/        (recorded interviews, calls)')
p.add_run('\n        - video/        (screen recordings, footage)')
p.add_run('\n        - metadata.json (incident metadata and file index)')
p.add_run('\n        - hashes.json   (SHA-256 hashes for each file)')

add_paragraph('\nExample: evidence_repository/INC-20260530-P5LZ/images/shot1.png')

# 5. User tutorial
add_heading('5. User Tutorial', level=1)
add_paragraph('Follow these steps to use FAFO:')
steps = [
    'Open the FAFO web application.',
    'Enter the correct password.',
    'Choose the correct role area: submitter, reviewer, lawyer, or admin.',
    'To submit a case, paste the social media source URL.',
    'Enter incident details, including category, target, notes, and submitter name.',
    'Upload screenshots, screen recordings, audio files, PDFs, or other evidence.',
    'Click “Submit Incident.”',
    'The system automatically creates an incident ID, timestamps the submission, stores uploaded files, generates file hashes, and sends an email notification.',
    'Reviewers can search cases, review evidence, and update case status.',
    'Admins can approve evidence for lawyer access.',
    'Lawyers can access approved evidence through the lawyer portal and download case packets.'
]
for i, s in enumerate(steps, start=1):
    add_paragraph(f'Step {i}: {s}')

# 6. Reflection paper
add_heading('6. Reflection Paper', level=1)
add_paragraph('This section reflects on the FAFO system, its societal impact, technical design, and lessons learned.')
add_paragraph('\nSummary of purpose and findings:\nFAFO demonstrates how accessible tools can support responsible digital evidence collection and preservation while balancing privacy and integrity. The modular architecture simplifies testing, extension, and replacement of components such as OCR or media analysis. Emphasis on role-based access ensures limited exposure of sensitive content to authorized personnel only.')
add_paragraph('\nTechnical lessons:\n- Runtime detection for system binaries (Tesseract/FFmpeg) increases portability.\n- Separating metadata from binary storage enables efficient indexing and secure transfer.\n- Cryptographic hashing and audit logging are critical for chain-of-custody.')

# Reflection Questions and Answers
add_heading('Reflection Questions', level=1)
qa = [
    ('Why is digital evidence preservation important?', 'Digital evidence preservation is crucial to maintain a verifiable record of events for investigations and legal proceedings. Without proper preservation, evidence may be altered, lost, or challenged in court.'),
    ('How does hashing support file integrity?', 'Hashing (e.g., SHA-256) provides a fingerprint of file content. Any modification changes the hash, enabling detection of tampering and ensuring integrity during transfer and storage.'),
    ('What privacy risks exist when storing sensitive evidence?', 'Risks include unauthorized access, data leakage, inadequate retention policies, and improper sharing. Sensitive metadata (names, locations) can expose individuals if not handled correctly.'),
    ('Why should the system remain private instead of public?', 'Keeping the system private minimizes exposure, reduces attack surface, and helps ensure controlled access to sensitive evidence. Public systems risk unauthorized retrieval or misuse.'),
    ('How can notification automation improve response time?', 'Automated notifications alert stakeholders immediately when new evidence is submitted, enabling faster triage and timely investigative actions.'),
    ('Why is role-based access important?', 'Role-based access enforces least privilege: users only see and perform actions allowed by their role, protecting sensitive data and maintaining separation of duties.'),
    ('What are the limitations of this system?', 'Limitations include reliance on host binaries for advanced processing, possible scalability constraints for large-scale media repositories, and the need for secure hosting and rigorous access controls.'),
    ('How could the system be improved in the future?', 'Future improvements: cloud-native object storage integration, strong encryption at rest, multi-factor authentication, immutable storage/backups, advanced analytics, and finer-grained access controls.')
]
for q, a in qa:
    add_paragraph(f'Q: {q}')
    add_paragraph(f'A: {a}')

# 7. Cost estimate
add_heading('7. Cost Estimate', level=1)
add_paragraph('High-level estimated monthly costs for a small production deployment (USD):')

# Add Table-like text
add_paragraph('Hosting & compute: $20 - $100 (small VM or managed app services)')
add_paragraph('Object storage (S3 or equivalent): $5 - $50 depending on volume')
add_paragraph('Database (managed PostgreSQL): $15 - $50')
add_paragraph('Backup & retention: $5 - $30')
add_paragraph('Domain, email, operational costs: $5 - $20')
add_paragraph('\nEstimated monthly total: $50 - $250 (varies with usage)')

# 8. Security/privacy explanation
add_heading('8. Security & Privacy Explanation', level=1)
add_paragraph('Security controls and recommendations:\n- Authentication: Enforce strong passwords and consider MFA.\n- Authorization: Role-based access control with least privilege.\n- Encryption: Use TLS in transit and encrypt sensitive data at rest.\n- Audit: Maintain immutable audit logs for all access and exports.\n- Hashing: Record SHA-256 hashes for integrity checks.\n- Retention: Define retention policies and secure deletion procedures.\n- Access approvals: Admin approval flow before granting lawyers access to sensitive packs.')

# Final project summary
add_heading('Final Project Summary', level=1)
add_paragraph('FAFO is a secure multimedia evidence preservation platform designed for documenting online incidents. It combines information systems, cybersecurity, digital forensics, automation, file management, and user workflow design. The system demonstrates how technology can support responsible documentation, evidence integrity, privacy, and organized incident response.')

# Footer
add_paragraph('\nDocument generated programmatically. Replace placeholders and insert screenshots as needed.')

# Save document
doc.save('FAFO_Documentation.docx')
print('FAFO_Documentation.docx created')
